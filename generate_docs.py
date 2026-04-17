import os
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Mendownload dependensi python-docx terlebih dahulu...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Dokumentasi Modul & File Sistem Informasi Manajemen Panti Asuhan (SIMPA)', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Dokumen ini berisi penjelasan mengenai struktur file dan fungsionalitas modul-modul yang ada dalam sistem SIMPA untuk mempermudah pemahaman dan pengembangan/perubahan sistem lebih lanjut.')

def add_module(name, desc, files):
    doc.add_heading(name, level=1)
    doc.add_paragraph(desc)
    
    for f_cat, f_list in files.items():
        p = doc.add_paragraph()
        p.add_run(f_cat + ':').bold = True
        
        for f in f_list:
            doc.add_paragraph(f, style='List Bullet')

modules = [
    {
        "name": "1. Modul Autentikasi & Pengguna (User Management)",
        "desc": "Modul ini menangani proses login, logout, dan manajemen data pengguna (Admin, Ketua, Bendahara, Donatur).",
        "files": {
            "Controllers (Logika)": ["AuthController.php (Menangani proses login/logout)", "UserController.php (Menangani proses CRUD data pengguna)"],
            "Models (Database)": ["User.php (Representasi struktur tabel users)"],
            "Views (Tampilan)": ["auth/ (Folder yang berisi tampilan form login)", "users/ (Folder tampilan tabel data pengguna dan form pengisian)"]
        }
    },
    {
        "name": "2. Modul Dashboard",
        "desc": "Modul ini berfungsi sebagai halaman utama setelah pengguna berhasil login, difungsikan untuk menampilkan ringkasan data statistik sistem (total anak asuh, total donasi, dll).",
        "files": {
            "Controllers (Logika)": ["DashboardController.php (Mengambil atau menyiapkan data data ringkasan untuk dashboard)"],
            "Views (Tampilan)": ["admin/dashboard.blade.php (Tampilan ringkasan statistik untuk Admin/Pengurus)", "donatur/dashboard.blade.php (Tampilan dashboard untuk Donatur)"]
        }
    },
    {
        "name": "3. Modul Manajemen Anak Asuh",
        "desc": "Modul utama untuk mendata dan mengelola biodata anak asuh yang berada di panti, termasuk status aktif/lulus dari panti.",
        "files": {
            "Controllers (Logika)": ["AnakAsuhController.php (Menangani proses tambah, ubah, dan hapus data)"],
            "Models (Database)": ["AnakAsuh.php (Model untuk tabel anak asuh)"],
            "Views (Tampilan)": ["anak-asuh/ (Folder yang berisi tampilan tabel daftar anak asuh dan formulir datanya)"]
        }
    },
    {
        "name": "4. Modul Donasi",
        "desc": "Modul sentral yang menangani pencatatan penerimaan donasi, baik dari input manual admin maupun verifikasi donasi dari halaman publik.",
        "files": {
            "Controllers (Logika)": ["DonasiController.php (Menangani pengelolaan seluruh data donasi beserta validasinya)"],
            "Models (Database)": ["Donasi.php (Model dari tabel donasi)"],
            "Views (Tampilan)": ["donasi/ (Folder yang memuat rekap riwayat serta verifikasi donasi masuk)"]
        }
    },
    {
        "name": "5. Modul Profil Donatur",
        "desc": "Modul khusus untuk donatur yang sudah punya akun guna melihat aktivitas/riwayat dukungan mereka serta mengelola profil pribadi.",
        "files": {
            "Controllers (Logika)": ["DonaturController.php (Menangani request halaman profil khusus donatur)"],
            "Models (Database)": ["Donatur.php (Model untuk profil spesifik milik donatur)"],
            "Views (Tampilan)": ["donatur/ (Folder yang berisi form profil donatur dan riwayat aktivitas mereka)"]
        }
    },
    {
        "name": "6. Modul Stok Panti (Logistik Gudang)",
        "desc": "Modul ini digunakan untuk memantau keluar masuk dan memperbarui stok barang habis pakai (contohnya beras, sembako, sabun) di panti asuhan.",
        "files": {
            "Controllers (Logika)": ["StokController.php (Mengelola sirkulasi dan status data stok barang)"],
            "Models (Database)": ["StokPanti.php (Model referensi ke tabel stok panti)"],
            "Views (Tampilan)": ["stok/ (Folder visualisasi pencatatan masuk-keluar stok gudang)"]
        }
    },
    {
        "name": "7. Modul Inventaris Peralatan",
        "desc": "Sistem pencatatan fasilitas, aset dan barang yang merupakan hak milik permanen (tidak habis pakai) panti asuhan, misal komputer, meja, kendaraan.",
        "files": {
            "Controllers (Logika)": ["InventarisController.php (Mengatur manajemen aset inventaris)"],
            "Models (Database)": ["InventarisPeralatan.php (Model referensi ke inventaris di db)"],
            "Views (Tampilan)": ["inventaris/ (Folder pencatatan pengadaan kondisi barang dan fasilitas)"]
        }
    },
    {
        "name": "8. Modul Perpustakaan & Peminjaman Buku",
        "desc": "Sistem mini manajemen koleksi judul buku bacaan atau edukasi di panti yang mencakup pencatatan rincian riwayat peminjaman oleh setiap anak asuh.",
        "files": {
            "Controllers (Logika)": ["PerpustakaanController.php (Mengelola pendataan katalog buku serta transaksi yang bersangkutan)"],
            "Models (Database)": ["Perpustakaan.php (Katalog utama buku)", "PeminjamanBuku.php (Model riwayat transaksi peminjaman)"],
            "Views (Tampilan)": ["perpustakaan/ (Folder yang membundel form pembuatan buku fisik dan form cek siapa yang baca)"]
        }
    },
    {
        "name": "9. Modul Pengeluaran Operasional",
        "desc": "Modul penunjang pembukuan finansial untuk merinci laporan keuangan keluar (baik biaya makan harian, utilitas, alat tulis).",
        "files": {
            "Controllers (Logika)": ["PengeluaranController.php (Mengurus form alur kas keluar operasi panti)"],
            "Models (Database)": ["Pengeluaran.php (Model laporan pengeluaran utilitas harian)"],
            "Views (Tampilan)": ["pengeluaran/ (Folder berisi rekapan histori keluar kas serta formulirnya)"]
        }
    },
    {
        "name": "10. Halaman Publik (UI Landing Page)",
        "desc": "Gerbang akses bagi donatur anonim atau khalayak masyarakat luar yang mengakses website. Digunakan untuk informasi sosialisasi atau formulir konfirmasi transfer anonim.",
        "files": {
            "Routes (Penghubung)": ["web.php (Di dalamnya terdefinisi routing '/' menuju homepage dan '/donasi-publik')"],
            "Views (Tampilan)": ["landing.blade.php (Muka/Beranda sistem ketika dibuka pengunjung lazim)", "welcome.blade.php (Bisa berupa template bawaan asli / alternate)"]
        }
    }
]

for mod in modules:
    add_module(mod['name'], mod['desc'], mod['files'])

doc.add_heading('Konfigurasi Master & Struktur Inti Laravel Utama', level=1)
doc.add_paragraph('Berikut beberapa file inti Laravel yang penting untuk diketahui jika ingin mengubah konfigurasi mendasar pada pengerjaan sistem di kemudian hari:', style='List Bullet')
core_files = [
    'routes/web.php : File pusat tempat mendefinisikan tautan (URL/Routes) aplikasi yang menunjuk ke Controller terkait.',
    '.env : File konfigurasi database (username, password), environment, nama sistem, email, dsb.',
    'app/Providers/ : Folder pengaturan startup/booting aplikasi. RouteServiceProvider dll.',
    'database/migrations/ : Folder esensial berisi cetak biru riwayat tabel/kolom yang dibentuk oleh database saat fresh install.',
    'resources/views/layouts/ : Folder yang menampung "Bungkus Utama Tampilan", seperti (Navbar, Footer, Sidebar admin) sehingga jika ada perubahan style menu cukup ubah file dilingkungan ini.'
]
for cf in core_files:
    doc.add_paragraph(cf, style='List Bullet')

doc.save('Dokumentasi_Modul_SIMPA.docx')
print("✅ File Dokumentasi_Modul_SIMPA.docx berhasil dibuat!")
